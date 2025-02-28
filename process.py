import os
import json
import docx
from typing import Dict, List, Optional
from langchain.text_splitter import RecursiveCharacterTextSplitter
from llama_index.core import Document, VectorStoreIndex, ServiceContext
# Fix for the import path - updated to the correct location
from llama_index.core.embeddings import HuggingFaceEmbedding
from pydantic import BaseModel, Field
from huggingface_hub import InferenceClient
from langchain_community.document_loaders import PyPDFLoader, TextLoader


from dotenv import load_dotenv
# Securely retrieve Hugging Face Token (avoid hardcoding in production)
load_dotenv()
TOKEN = os.getenv("HF_TOKEN")  # Ensure you set this in your environment variables
MODEL_ID = "mistralai/Mistral-7B-Instruct-v0.3"

# Initialize Hugging Face Inference Client
client = InferenceClient(token=TOKEN)

# Define Pydantic model for topics and subtopics
class Topic(BaseModel):
    topic: str = Field(description="Main topic name")
    subtopics: List[str] = Field(description="List of subtopics under the main topic")

# Predefined topics
TOPICS = [
    Topic(
        topic="Circulatory System",
        subtopics=[
            "Anatomy of the Heart",
            "Blood Circulation",
            "Blood Vessels and Their Functions",
            "Blood Pressure Regulation",
            "Cardiac Cycle and Conduction System",
            "Hemodynamics",
            "Common Circulatory Diseases",
            "Role of the Circulatory System in Homeostasis",
        ],
    )
]

# Custom function to extract text from docx files
def extract_text_from_docx(file_path: str) -> str:
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Extract text from PDF files
def extract_text_from_pdf(file_path: str) -> str:
    loader = PyPDFLoader(file_path)
    return "\n".join([page.page_content for page in loader.load()])

# Extract text from text files
def extract_text_from_txt(file_path: str) -> str:
    loader = TextLoader(file_path)
    return "\n".join([page.page_content for page in loader.load()])

# Extract text from files based on extension
def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.docx':
            return extract_text_from_docx(file_path)
        elif ext == '.pdf':
            return extract_text_from_pdf(file_path)
        elif ext == '.txt':
            return extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from {file_path}: {e}")

# Load documents from a folder
def load_documents(folder_path: str) -> List[Dict[str, str]]:
    if not os.path.exists(folder_path):
        raise ValueError(f"Folder {folder_path} does not exist")
    
    documents = []
    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        if os.path.isfile(file_path):
            try:
                text = extract_text(file_path)
                documents.append({"filename": file_name, "content": text})
            except Exception as e:
                print(f"Error processing {file_name}: {e}")
    return documents

# Hugging Face LLM Wrapper
class HuggingFaceInferenceLLM:
    def __init__(self, model_id: str = MODEL_ID, token: Optional[str] = TOKEN):
        self.client = InferenceClient(token=token)
        self.model_id = model_id

    def __call__(self, text: str, query: Optional[str] = None) -> str:
        prompt = f"Instruction: {query}\n\nContent: {text}" if query else text
        try:
            response = self.client.text_generation(
                prompt,
                model=self.model_id,
                max_new_tokens=200,
                temperature=0.7,
                top_p=0.9,
                do_sample=True,  # Enable sampling for better responses
            )
            return response
        except Exception as e:
            return f"Error processing query with Inference API: {e}"

# Analyze documents
def analyze_documents_by_subtopics(documents: List[Dict[str, str]], topics: List[Topic]) -> Dict[str, Dict[str, str]]:
    llm = HuggingFaceInferenceLLM()
    
    # Convert to Llama-Index Document objects
    llama_docs = [Document(text=doc["content"], metadata={"filename": doc["filename"]}) for doc in documents]
    
    # Split text into manageable chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = [Document(text=chunk, metadata={"filename": doc.metadata["filename"]}) 
                  for doc in llama_docs for chunk in text_splitter.split_text(doc.text)]
    
    # Use HuggingFace embeddings instead of OpenAI
    embed_model = HuggingFaceEmbedding(model_name="sentence-transformers/all-MiniLM-L6-v2")
    service_context = ServiceContext.from_defaults(embed_model=embed_model, llm=llm)
    
    # Create index and query engine
    index = VectorStoreIndex.from_documents(split_docs, service_context=service_context)
    query_engine = index.as_query_engine()
    
    # Structure results
    structured_results: Dict[str, Dict[str, str]] = {}
    for topic_entry in topics:
        topic = topic_entry.topic
        structured_results[topic] = {}
        for subtopic in topic_entry.subtopics:
            query = f"Extract key information and summarize content related to '{subtopic}'"
            response = query_engine.query(query)
            structured_results[topic][subtopic] = str(response)
    
    return structured_results

def main() -> None:
    folder_path = "./content_materials"  # Ensure this path exists
    try:
        documents = load_documents(folder_path)
        if not documents:
            print("No valid documents found to process.")
        else:
            results = analyze_documents_by_subtopics(documents, TOPICS)
            with open("subtopic_analysis_results.json", "w", encoding="utf-8") as f:
                json.dump(results, f, indent=4, ensure_ascii=False)
            print("Analysis complete. Results saved to subtopic_analysis_results.json")
    except Exception as e:
        print(f"Pipeline failed: {e}")

if __name__ == "__main__":
    main()